import io
import markdown
from docx import Document
from docx.shared import Pt, Inches
from xhtml2pdf import pisa


async def convert_markdown_to_pdf(markdown_content: str) -> bytes:
    """
    Convert markdown content to PDF.
    
    Args:
        markdown_content: Markdown content to convert
    
    Returns:
        PDF bytes
    """
    # Convert markdown to HTML
    html = markdown.markdown(markdown_content)
    
    # Add some basic styling
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.4;
                margin: 0.5in;
            }}
            h1 {{
                font-size: 16pt;
                margin-bottom: 0.2in;
            }}
            h2 {{
                font-size: 14pt;
                margin-top: 0.2in;
                margin-bottom: 0.1in;
            }}
            h3 {{
                font-size: 12pt;
                margin-top: 0.1in;
                margin-bottom: 0.1in;
            }}
            p {{
                margin-bottom: 0.1in;
            }}
            ul {{
                margin-top: 0.05in;
                margin-bottom: 0.1in;
            }}
            li {{
                margin-bottom: 0.05in;
            }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    
    # Convert HTML to PDF
    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(
        src=styled_html,
        dest=pdf_buffer
    )
    
    if pisa_status.err:
        raise Exception("Error converting HTML to PDF")
    
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()


async def convert_markdown_to_docx(markdown_content: str) -> bytes:
    """
    Convert markdown content to DOCX.
    
    Args:
        markdown_content: Markdown content to convert
    
    Returns:
        DOCX bytes
    """
    # Create a new Document
    doc = Document()
    
    # Set up margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content by lines for basic parsing
    lines = markdown_content.split('\n')
    
    in_list = False
    list_items = []
    
    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for blank lines
            if in_list:
                # End the current list
                for item in list_items:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(item)
                list_items = []
                in_list = False
            
            doc.add_paragraph()
            continue
        
        # Handle headings
        if line.startswith('# '):
            if in_list:
                # End the current list
                for item in list_items:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(item)
                list_items = []
                in_list = False
            
            p = doc.add_heading(line[2:], 0)  # Heading 1
        elif line.startswith('## '):
            if in_list:
                # End the current list
                for item in list_items:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(item)
                list_items = []
                in_list = False
            
            p = doc.add_heading(line[3:], 1)  # Heading 2
        elif line.startswith('### '):
            if in_list:
                # End the current list
                for item in list_items:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(item)
                list_items = []
                in_list = False
            
            p = doc.add_heading(line[4:], 2)  # Heading 3
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            in_list = True
            list_items.append(line[2:])
        # Regular paragraph
        else:
            if in_list:
                # End the current list
                for item in list_items:
                    p = doc.add_paragraph(style='ListBullet')
                    p.add_run(item)
                list_items = []
                in_list = False
            
            # Handle bold and italic inline formatting
            p = doc.add_paragraph()
            current_text = ""
            is_bold = False
            is_italic = False
            
            i = 0
            while i < len(line):
                if i+1 < len(line) and line[i:i+2] == '**' and not is_italic:
                    if current_text:
                        run = p.add_run(current_text)
                        run.bold = is_bold
                        current_text = ""
                    is_bold = not is_bold
                    i += 2
                elif i+1 < len(line) and line[i:i+2] == '__' and not is_italic:
                    if current_text:
                        run = p.add_run(current_text)
                        run.bold = is_bold
                        current_text = ""
                    is_bold = not is_bold
                    i += 2
                elif line[i] == '*' and not is_bold:
                    if current_text:
                        run = p.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        current_text = ""
                    is_italic = not is_italic
                    i += 1
                elif line[i] == '_' and not is_bold:
                    if current_text:
                        run = p.add_run(current_text)
                        run.bold = is_bold
                        run.italic = is_italic
                        current_text = ""
                    is_italic = not is_italic
                    i += 1
                else:
                    current_text += line[i]
                    i += 1
            
            if current_text:
                run = p.add_run(current_text)
                run.bold = is_bold
                run.italic = is_italic
    
    # Add any remaining list items
    if in_list:
        for item in list_items:
            p = doc.add_paragraph(style='ListBullet')
            p.add_run(item)
    
    # Save the document to a BytesIO object
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue()

import os
try:
    from docxtpl import DocxTemplate
except Exception:  # pragma: no cover - optional dependency
    DocxTemplate = None  # type: ignore


class TemplateProcessor:
    """Handles docx template processing for resume customization."""

    def __init__(self, templates_dir: str = "resume_templates") -> None:
        self.templates_dir = templates_dir

    async def get_template(self, template_id: str) -> str:
        """Get a template by its ID."""
        template_path = os.path.join(self.templates_dir, f"{template_id}.docx")
        if not os.path.exists(template_path):
            raise ValueError(f"Template {template_id} not found")
        return template_path

    async def apply_template(self, template_id: str, context_data: dict) -> io.BytesIO:
        """Apply customized content to a template."""
        template_path = await self.get_template(template_id)
        if DocxTemplate is None:
            raise ImportError("docxtpl package is required for template processing")

        doc = DocxTemplate(template_path)
        doc.render(context_data)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    async def parse_resume_to_context(self, resume_content: str, structure: dict | None = None) -> dict:
        """Parse resume content into context data for templates."""
        sections: dict[str, list[str]] = {}
        current_section = "header"
        sections[current_section] = []

        for line in resume_content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.isupper() or line.endswith(":"):
                current_section = line.lower().replace(":", "").strip()
                sections[current_section] = []
            else:
                sections[current_section].append(line)

        context = {sec: "\n".join(lines) for sec, lines in sections.items()}
        return context
